#!/usr/bin/env python3
"""
Agent Matrix — 工具注册中心 (Tool Registry)
===========================================
为子 Agent 提供原生 function calling 可调用的工具集。

设计原则：
  - 首批仅内置「只读、安全」工具，不给写库/删除类能力。
  - 工具执行统一带 try/except 兜底，失败返回字符串错误信息而非抛异常，
    保证 ReAct 循环不会因单个工具出错而崩溃。
  - 按 Agent 的 allowed_tools 白名单过滤，未授权工具不下发给模型。

对外接口：
  - get_tools_for_agent(allowed_tools) -> list[schema]
  - execute_tool(name, args) -> str
"""
from i18n import _
import json, os, sys, logging
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..'))

logger = logging.getLogger(__name__)

# ============================================================
# 工具 Schema（OpenAI function calling 格式）
# ============================================================

TOOL_SCHEMAS = {
    "get_system_health": {
        "type": "function",
        "function": {
            "name": "get_system_health",
            "description": "获取系统最近一次健康巡检的结果汇总，包括健康分、通过/警告/错误数量，以及各检查项状态。用于回答系统运行状态、服务健康、告警相关问题。",
            "parameters": {"type": "object", "properties": {}, "required": []}
        }
    },
    "query_stats": {
        "type": "function",
        "function": {
            "name": "query_stats",
            "description": "查询站点访问统计报告（PV/UV/会话/趋势/来源/热门页面），返回可读的文字洞察。用于回答流量、访问量、数据趋势相关问题。",
            "parameters": {
                "type": "object",
                "properties": {
                    "days": {
                        "type": "integer",
                        "description": "统计周期天数，默认 7 天",
                        "default": 7
                    }
                },
                "required": []
            }
        }
    },
    "search_knowledge": {
        "type": "function",
        "function": {
            "name": "search_knowledge",
            "description": "在平台知识库中检索与关键词相关的内容片段。用于回答产品功能、FAQ、使用帮助相关问题。",
            "parameters": {
                "type": "object",
                "properties": {
                    "keyword": {
                        "type": "string",
                        "description": "检索关键词"
                    }
                },
                "required": ["keyword"]
            }
        }
    },
    "ads_list": {
        "type": "function",
        "function": {
            "name": "ads_list",
            "description": "列出广告管理系统中的广告位，可按站点、位置、是否启用筛选。",
            "parameters": {
                "type": "object",
                "properties": {
                    "site_key": {"type": "string", "description": "站点标识，默认 default"},
                    "position": {"type": "string", "description": "广告位置，例如 sidebar"},
                    "active_only": {"type": "boolean", "description": "是否只返回启用状态的广告", "default": False}
                },
                "required": []
            }
        }
    },
    "ads_create": {
        "type": "function",
        "function": {
            "name": "ads_create",
            "description": "创建一个新的广告位。支持图片广告或广告代码，可设置投放位置、时间、定向规则、权重等。",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "广告名称（必填）"},
                    "site_key": {"type": "string", "description": "站点标识，默认 default"},
                    "zone_id": {"type": "integer", "description": "广告区域 ID，默认 0"},
                    "position": {"type": "string", "description": "广告位置标识，默认 sidebar"},
                    "page": {"type": "string", "description": "展示页面，默认 * 表示全站"},
                    "ad_type": {"type": "string", "enum": ["image", "code"], "description": "广告类型"},
                    "image_url": {"type": "string", "description": "图片广告 URL"},
                    "link_url": {"type": "string", "description": "图片广告跳转链接"},
                    "ad_code": {"type": "string", "description": "广告代码（HTML/JS）"},
                    "width": {"type": "integer", "description": "宽度（px）"},
                    "height": {"type": "integer", "description": "高度（px）"},
                    "targeting_rules": {"type": "object", "description": "定向规则 JSON 对象"},
                    "schedule_start": {"type": "string", "description": "投放开始时间 ISO 格式"},
                    "schedule_end": {"type": "string", "description": "投放结束时间 ISO 格式"},
                    "weight": {"type": "integer", "description": "权重，默认 1"},
                    "freq_cap": {"type": "integer", "description": "每用户每日频次上限，0 表示无限制"},
                    "click_tag": {"type": "string", "description": "点击追踪标记"},
                    "utm_source": {"type": "string", "description": "UTM 来源"},
                    "is_active": {"type": "integer", "description": "是否启用，1=启用，0=禁用"},
                    "sort_order": {"type": "integer", "description": "排序，默认 0"}
                },
                "required": ["name"]
            }
        }
    },
    "ads_update": {
        "type": "function",
        "function": {
            "name": "ads_update",
            "description": "更新指定广告位的字段，例如启用/禁用、修改代码、调整权重等。",
            "parameters": {
                "type": "object",
                "properties": {
                    "ad_id": {"type": "integer", "description": "广告 ID"},
                    "updates": {"type": "object", "description": "要更新的字段键值对，字段含义同 ads_create"}
                },
                "required": ["ad_id", "updates"]
            }
        }
    },
    "ads_delete": {
        "type": "function",
        "function": {
            "name": "ads_delete",
            "description": "删除指定广告位。",
            "parameters": {
                "type": "object",
                "properties": {
                    "ad_id": {"type": "integer", "description": "广告 ID"}
                },
                "required": ["ad_id"]
            }
        }
    },
    "ads_get_stats": {
        "type": "function",
        "function": {
            "name": "ads_get_stats",
            "description": "查询广告统计数据，包括展示量、点击量、CTR 及每日趋势。",
            "parameters": {
                "type": "object",
                "properties": {
                    "ad_id": {"type": "integer", "description": "广告 ID，不传则统计全部广告"},
                    "days": {"type": "integer", "description": "统计天数，默认 7", "default": 7}
                },
                "required": []
            }
        }
    },
    "ads_analyze": {
        "type": "function",
        "function": {
            "name": "ads_analyze",
            "description": "分析广告效果，返回高点击、低 CTR、趋势等文字洞察与优化建议。",
            "parameters": {
                "type": "object",
                "properties": {
                    "days": {"type": "integer", "description": "统计天数，默认 7", "default": 7}
                },
                "required": []
            }
        }
    },
    "ads_render_snippet": {
        "type": "function",
        "function": {
            "name": "ads_render_snippet",
            "description": "生成一段 Jinja2 模板代码，用于在页面指定位置渲染广告位。",
            "parameters": {
                "type": "object",
                "properties": {
                    "position": {"type": "string", "description": "广告位置标识，默认 sidebar"},
                    "page": {"type": "string", "description": "展示页面，默认 *"},
                    "site_key": {"type": "string", "description": "站点标识，默认 default"},
                    "zone_id": {"type": "integer", "description": "广告区域 ID"}
                },
                "required": ["position"]
            }
        }
    },
    "generate_ppt": {
        "type": "function",
        "function": {
            "name": "generate_ppt",
            "description": "使用 AI 生成 PowerPoint 演示文稿（PPTX）。用户提供主题、页数和风格即可生成一份可直接下载的 PPT 文件。",
            "parameters": {
                "type": "object",
                "properties": {
                    "topic": {"type": "string", "description": "PPT 主题（必填）"},
                    "pages": {"type": "integer", "description": "页数，默认 10，范围 3-20", "default": 10},
                    "style": {"type": "string", "description": "风格描述，如'Dark 科技风'、'简约商务'、'教育风格'等", "default": _("Dark Tech Style, 16:9")}
                },
                "required": ["topic"]
            }
        }
    },
    "generate_image": {
        "type": "function",
        "function": {
            "name": "generate_image",
            "description": "使用 AI 生成图像。用户提供描述文字即可生成图片，支持指定风格和数量。",
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {"type": "string", "description": "图像描述（必填）"},
                    "style": {"type": "string", "description": "风格，可选 realistic/anime/3d/painting/line-art", "default": "realistic"},
                    "count": {"type": "integer", "description": "生成数量（1-4），默认 1", "default": 1},
                    "size": {"type": "string", "description": "尺寸，可选 1024x1024/1792x1024/1024x1792，默认 1024x1024", "default": "1024x1024"}
                },
                "required": ["prompt"]
            }
        }
    },
    "generate_markdown": {
        "type": "function",
        "function": {
            "name": "generate_markdown",
            "description": "使用 AI 生成 Markdown 文档。根据用户提供主题和格式要求，生成可直接预览或下载的 .md 文件。支持文章、报告、技术文档、笔记等。",
            "parameters": {
                "type": "object",
                "properties": {
                    "topic": {"type": "string", "description": "文档主题（必填）"},
                    "outline": {"type": "string", "description": "大纲要求，可选章节或要点"},
                    "style": {"type": "string", "description": "风格，可选 technical/professional/creative/simple/list", "default": "technical"}
                },
                "required": ["topic"]
            }
        }
    },
    "generate_docx": {
        "type": "function",
        "function": {
            "name": "generate_docx",
            "description": "使用 AI 生成 Word 文档（.docx）。根据用户提供主题和格式要求，生成样式化的 Word 文档，包含标题、段落、列表、表格等。适合生成报告、合同、方案书、简历等正式文档。",
            "parameters": {
                "type": "object",
                "properties": {
                    "topic": {"type": "string", "description": "文档主题（必填）"},
                    "style": {"type": "string", "description": "风格，可选 professional/formal/creative/simple/report", "default": "professional"},
                    "sections": {"type": "integer", "description": "章节数量（3-12），默认 5", "default": 5}
                },
                "required": ["topic"]
            }
        }
    },
    "cms_list_channels": {
        "type": "function",
        "function": {
            "name": "cms_list_channels",
            "description": "列出 CMS 系统中所有可用的发布频道（分类/栏目），用于确定文章可以发布到哪些频道。",
            "parameters": {"type": "object", "properties": {}, "required": []}
        }
    },
    "cms_create_post": {
        "type": "function",
        "function": {
            "name": "cms_create_post",
            "description": "在 CMS 系统中创建一篇新文章。创建后默认为草稿状态，需要调用 cms_publish_post 发布。",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "文章标题（必填）"},
                    "content": {"type": "string", "description": "文章正文内容，支持 HTML 格式"},
                    "excerpt": {"type": "string", "description": "文章摘要/导语"},
                    "category": {"type": "string", "description": "发布频道/分类标识，如 insights、news"},
                    "cover_image": {"type": "string", "description": "封面图片 URL"},
                    "tags": {"type": "array", "items": {"type": "string"}, "description": "标签列表"}
                },
                "required": ["title"]
            }
        }
    },
    "cms_publish_post": {
        "type": "function",
        "function": {
            "name": "cms_publish_post",
            "description": "将 CMS 文章发布到指定频道。可选择本地频道和社交媒体平台。",
            "parameters": {
                "type": "object",
                "properties": {
                    "post_id": {"type": "integer", "description": "文章 ID（必填）"},
                    "channels": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "发布目标频道列表，格式为 'local:频道名'，如 ['local:insights']"
                    }
                },
                "required": ["post_id", "channels"]
            }
        }
    },
}


# ============================================================
# 工具执行器（均为只读）
# ============================================================

def _get_matrix_db():
    """获取 agent_matrix 所在主库连接（复用项目 models.get_db）"""
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
    from models import get_db
    return get_db()


def _tool_get_system_health(args):
    """读取最近一次健康巡检结果汇总（只读）"""
    try:
        sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
        from plugins.health_check.models import get_db as health_db
        with health_db() as conn:
            run = conn.execute(
                "SELECT * FROM check_runs WHERE status='completed' "
                "ORDER BY created_at DESC LIMIT 1"
            ).fetchone()
            if not run:
                return _("No health inspection records.")
            run = dict(run)
            total = run.get('total_checks', 0) or 0
            passed = run.get('passed', 0) or 0
            warnings = run.get('warnings', 0) or 0
            errors = run.get('errors', 0) or 0
            denom = passed + warnings + errors
            score = round((passed + warnings * 0.5) * 100 / denom, 1) if denom else 100.0

            items = conn.execute(
                "SELECT check_name, status, message FROM check_history "
                "WHERE run_id=%s ORDER BY status DESC",
                (run['id'],)
            ).fetchall()
        lines = [
            f"Health Score: {score}/100",
            f"Total Checks: {total}, Passed: {passed}, Warnings: {warnings}, Errors: {errors}",
            f"Inspection Time: {run.get('created_at', '')}",
        ]
        abnormal = [dict(i) for i in items if i['status'] != 'passed']
        if abnormal:
            lines.append(_("Abnormal items:"))
            for i in abnormal[:15]:
                lines.append(f"  - [{i['status']}] {i['check_name']}: {(i['message'] or '')[:80]}")
        else:
            lines.append(_("All checks passed."))
        return '\n'.join(lines)
    except Exception as e:
        logger.warning(f"[tool:get_system_health] 执行失败: {e}")
        return f"Failed to Get Health Status: {e}"


def _tool_query_stats(args):
    """生成站点统计报告的文字洞察（只读）"""
    try:
        days = int(args.get('days', 7) or 7)
        days = max(1, min(days, 90))
        sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
        from plugins.analytics.tracker import generate_report, generate_insight_text
        report = generate_report(days=days)
        return generate_insight_text(report)
    except Exception as e:
        logger.warning(f"[tool:query_stats] 执行失败: {e}")
        return f"Failed to Query Data Statistics: {e}"


def _tool_search_knowledge(args):
    """在知识库中检索关键词（只读）"""
    try:
        keyword = str(args.get('keyword', '')).strip()
        if not keyword:
            return _("No search keyword provided.")
        with _get_matrix_db() as conn:
            row = conn.execute(
                "SELECT value FROM system_config WHERE key='chatbot_knowledge_base'"
            ).fetchone()
        content = (row['value'] if row and row['value'] else '') or ''
        if not content:
            return _("The knowledge base is empty.")
        # 简单按段落匹配，返回命中片段
        blocks = [b.strip() for b in content.split('\n\n') if b.strip()]
        hits = [b for b in blocks if keyword.lower() in b.lower()]
        if not hits:
            return f"No content related to 「{keyword}」 found in the knowledge base."
        return '\n---\n'.join(hits[:5])[:2000]
    except Exception as e:
        logger.warning(f"[tool:search_knowledge] 执行失败: {e}")
        return f"Failed to retrieve knowledge base: {e}"


def _tool_ads_list(args):
    """列出广告"""
    try:
        import plugins.ads.ai_tools as ads_tools
        res = ads_tools.list_ads(
            site_key=args.get('site_key'),
            position=args.get('position'),
            active_only=args.get('active_only', False)
        )
        if not res['success']:
            return f"Failed to Get Ad List: {res.get('error')}"
        ads = res.get('data', [])
        if not ads:
            return _("No ad space.")
        lines = [f"Total {len(ads)} Ad Positions:"]
        for a in ads:
            status = _('Enable') if a.get('is_active') else _('Deactivate')
            lines.append(
                f"ID {a['id']}: {a['name']} | Site {a.get('site_key','default')} |"
                f"Location {a.get('position','-')} | Type {a.get('ad_type','image')} | {status}"
            )
        return '\n'.join(lines)
    except Exception as e:
        logger.warning(f"[tool:ads_list] 执行失败: {e}")
        return f"Failed to Get Ad List: {e}"


def _tool_ads_create(args):
    """创建广告"""
    try:
        import plugins.ads.ai_tools as ads_tools
        res = ads_tools.create_ad(args)
        if res['success']:
            return f"✅ Ad created, ID: {res['data']['id']}"
        return f"❌ Creation Failed: {res.get('error')}"
    except Exception as e:
        logger.warning(f"[tool:ads_create] 执行失败: {e}")
        return f"Ad creation failed: {e}"


def _tool_ads_update(args):
    """更新广告"""
    try:
        import plugins.ads.ai_tools as ads_tools
        ad_id = args.get('ad_id')
        updates = args.get('updates', {})
        res = ads_tools.update_ad(ad_id, updates)
        if res['success']:
            return f"✅ Ad {ad_id} updated"
        return f"❌ Update Failed: {res.get('error')}"
    except Exception as e:
        logger.warning(f"[tool:ads_update] 执行失败: {e}")
        return f"Failed to update ad: {e}"


def _tool_ads_delete(args):
    """删除广告"""
    try:
        import plugins.ads.ai_tools as ads_tools
        res = ads_tools.delete_ad(args.get('ad_id'))
        if res['success']:
            return f"✅ Ad {args.get('ad_id')} deleted"
        return f"❌ Deletion Failed: {res.get('error')}"
    except Exception as e:
        logger.warning(f"[tool:ads_delete] 执行失败: {e}")
        return f"Failed to Delete Ad: {e}"


def _tool_ads_get_stats(args):
    """查询广告统计"""
    try:
        import plugins.ads.ai_tools as ads_tools
        res = ads_tools.get_stats(
            ad_id=args.get('ad_id'),
            site_key=args.get('site_key'),
            days=int(args.get('days', 7))
        )
        if not res['success']:
            return f"Query Statistics Failed: {res.get('error')}"
        data = res.get('data', {})
        total = data.get('total', {})
        daily = data.get('daily', [])
        lines = [
            f"=== Advertising Statistics (Last {args.get('days',7)} Days) ===",
            f"Impressions: {total.get('impressions', 0)}",
            f"Clicks: {total.get('clicks', 0)}",
            f"CTR: {total.get('ctr', 0)}%",
        ]
        if daily:
            lines.append(_("Daily Trend:"))
            for r in daily[-10:]:
                lines.append(f"  {r['stat_date']}: Impressions {r.get('impressions',0)} Clicks {r.get('clicks',0)}")
        return '\n'.join(lines)
    except Exception as e:
        logger.warning(f"[tool:ads_get_stats] 执行失败: {e}")
        return f"Failed to Query Ad Statistics: {e}"


def _tool_ads_analyze(args):
    """分析广告效果"""
    try:
        import plugins.ads.ai_tools as ads_tools
        res = ads_tools.analyze_ads(days=int(args.get('days', 7)))
        if res['success']:
            return res['data']
        return f"Analysis failed: {res.get('error')}"
    except Exception as e:
        logger.warning(f"[tool:ads_analyze] 执行失败: {e}")
        return f"Advertisement analysis failed: {e}"


def _tool_ads_render_snippet(args):
    """生成广告渲染代码片段"""
    try:
        import plugins.ads.ai_tools as ads_tools
        res = ads_tools.generate_render_snippet(
            position=args.get('position', 'sidebar'),
            page=args.get('page', '*'),
            site_key=args.get('site_key', 'default'),
            zone_id=args.get('zone_id')
        )
        if res['success']:
            return "在模板中加入以下代码即可渲染广告位：\n```jinja2\n" + res['data'] + "\n```"
        return f"Failed to generate code: {res.get('error')}"
    except Exception as e:
        logger.warning(f"[tool:ads_render_snippet] 执行失败: {e}")
        return f"Failed to generate ad rendering code: {e}"


def _tool_cms_list_channels(args):
    """列出 CMS 频道"""
    try:
        from models import get_db
        with get_db() as conn:
            rows = conn.execute(
                "SELECT id, name, slug, description FROM cms_categories WHERE is_active=1 ORDER BY sort_order"
            ).fetchall()
        if not rows:
            return "No available channels."
        lines = ["Available CMS channels:"]
        for r in rows:
            r = dict(r)
            lines.append(f"  ID {r['id']}: {r['name']} (slug: {r['slug']})")
        return '\n'.join(lines)
    except Exception as e:
        logger.warning(f"[tool:cms_list_channels] Failed: {e}")
        return f"Failed to get channel list: {e}"


def _tool_cms_create_post(args):
    """创建 CMS 文章"""
    try:
        title = str(args.get('title', '')).strip()
        if not title:
            return "Error: Article title is required."

        import uuid
        post_data = {
            'title': title,
            'content': str(args.get('content', '')),
            'excerpt': str(args.get('excerpt', '')),
            'category': str(args.get('category', 'insights')),
            'cover_image': str(args.get('cover_image', '')),
            'tags': args.get('tags', []),
            'slug': 'article-' + str(uuid.uuid4())[:8],
            'source': 'ai',
            'is_published': 0,
        }

        sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'auth-center'))
        from models.cms import upsert_post
        result = upsert_post(post_data)
        post_id = result.get('id') if isinstance(result, dict) else result
        return f"Article created, ID: {post_id}, Title: {title}"
    except Exception as e:
        logger.warning(f"[tool:cms_create_post] Failed: {e}")
        return f"Failed to create article: {e}"


def _tool_cms_publish_post(args):
    """发布 CMS 文章"""
    try:
        post_id = args.get('post_id')
        channels = args.get('channels', [])

        if not post_id:
            return "Error: post_id is required."
        if not channels:
            return "Error: channels is required."

        from models import get_db
        with get_db() as conn:
            row = conn.execute(
                "SELECT * FROM cms_posts WHERE id=%s", (post_id,)
            ).fetchone()
        if not row:
            return f"Error: Article {post_id} not found."

        # Separate local channels
        local_cats = []
        for ch in channels:
            if ch.startswith('local:'):
                local_cats.append(ch.split(':', 1)[1])

        if local_cats:
            sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'auth-center'))
        from models.cms import upsert_post
            post = dict(row)
            upsert_post({
                'id': post_id,
                'slug': post['slug'],
                'category': local_cats[0],
                'title': post['title'],
                'excerpt': post.get('excerpt', ''),
                'content': post['content'],
                'is_published': 1,
                'publish_channels': channels,
            })

        return f"Article {post_id} published to: {', '.join(local_cats)}"
    except Exception as e:
        logger.warning(f"[tool:cms_publish_post] Failed: {e}")
        return f"Failed to publish article: {e}"


def _tool_generate_ppt(args):
    """使用 AI 生成 PPT 文件"""
    try:
        topic = str(args.get('topic', _('Untitled topic'))).strip()
        if not topic:
            return '❌ 请提供 PPT 主题'
        pages = max(3, min(int(args.get('pages', 10) or 10), 20))
        style = str(args.get('style', _('Dark Tech Style, 16:9')))
        from agent_matrix.routes import _generate_ppt_file
        filename = _generate_ppt_file(topic, pages, style)
        if filename:
            return f'✅ PPT 已生成："{topic}"（{pages}页）\n下载链接：/admin/agent-matrix/media/download/{filename}'
        return '❌ PPT 生成失败，请检查后端日志'
    except Exception as e:
        logger.warning(f"[tool:generate_ppt] 执行失败: {e}")
        return f'❌ PPT Generation Error: {e}'


def _tool_generate_image(args):
    """使用硅基流动 API 生成图像"""
    try:
        prompt = str(args.get('prompt', '')).strip()
        if not prompt:
            return '❌ 请提供图像描述'
        count = max(1, min(int(args.get('count', 1) or 1), 4))
        size = str(args.get('size', '1024x1024'))
        if '\\' in size:
            size = '1024x1024'

        # 从 system_config 读取 siliconflow_api_key
        from models import get_db
        with get_db() as conn:
            row = conn.execute("SELECT value FROM system_config WHERE key='siliconflow_api_key'").fetchone()
        api_key = row['value'] if row else os.environ.get('SILICONFLOW_API_KEY', '')
        if not api_key:
            return _('❌ Si Ji Liu Dong API Key Not Configured')

        # 风格 → 模型映射
        style_map = {
            'realistic': 'black-forest-labs/FLUX.1-dev',
            'anime': 'stabilityai/stable-diffusion-xl-base-1.0',
            '3d': 'stabilityai/stable-diffusion-xl-base-1.0',
            'painting': 'SG161222/RealVisXL_V4.0',
            'line-art': 'lykon/dreamshaper-xl-1-0',
        }
        model = style_map.get(args.get('style', 'realistic'), 'black-forest-labs/FLUX.1-dev')

        from openai import OpenAI
        client = OpenAI(api_key=api_key, base_url='https://api.siliconflow.cn/v1')
        resp = client.images.generate(
            model=model,
            prompt=prompt,
            n=count,
            size=size
        )

        # 保存图片到 media 目录
        import uuid, requests
        media_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'data', 'media', 'temp')
        os.makedirs(media_dir, exist_ok=True)

        urls = []
        for i, img_data in enumerate(resp.data):
            img_url = img_data.url
            if not img_url:
                continue
            r = requests.get(img_url, timeout=30)
            ext = 'png'
            fn = f"img_{uuid.uuid4().hex[:8]}_{i}.{ext}"
            fp = os.path.join(media_dir, fn)
            with open(fp, 'wb') as f:
                f.write(r.content)
            download_url = f'/admin/agent-matrix/media/download/{fn}'
            urls.append(download_url)

        if not urls:
            return _('❌ Image Generation API Returned Empty')

        lines = [f'✅ Generated {len(urls)} images:']
        for u in urls:
            lines.append(f'  {u}')
        return '\n'.join(lines)

    except Exception as e:
        logger.warning(f"[tool:generate_image] 执行失败: {e}")
        return f'❌ Image Generation Error: {e}'


def _tool_generate_markdown(args):
    """使用 AI 生成 Markdown 文档"""
    try:
        topic = str(args.get('topic', '')).strip()
        if not topic:
            return '❌ 请提供文档主题'
        outline = str(args.get('outline', ''))
        style = str(args.get('style', 'technical'))

        style_prompt = {
            'technical': _('Technical Document Style: Clear structure, well-defined hierarchy, professional terms in English'),
            'professional': _('Business report style, formal language, complete paragraphs, clear conclusions'),
            'creative': _('Creative Writing Style, Lively and Persuasive Language'),
            'simple': _('Concise style, bullet points for quick reading'),
            'list': _('List Style, Mainly Items and Sub-items'),
        }.get(style, _('Technical Document Style'))

        prompt_text = f'请撰写一篇关于"{topic}"的Markdown文档。\n风格要求：{style_prompt}'
        if outline:
            prompt_text += f'\n大纲要求：{outline}'
        prompt_text += '\n请直接输出Markdown格式内容，包含标题（#）、段落、列表、代码块等格式化元素。'

        # 读取 SiliconFlow API Key（走硅基流动默认）
        from models import get_db
        with get_db() as conn:
            row = conn.execute("SELECT value FROM system_config WHERE key='siliconflow_api_key'").fetchone()
        api_key = row['value'] if row else os.environ.get('SILICONFLOW_API_KEY', '')
        if not api_key:
            # 兜底 dashscope
            with get_db() as conn:
                row = conn.execute("SELECT value FROM system_config WHERE key='dashscope_text_key'").fetchone()
            if row:
                api_key = row['value']
                base_url = 'https://dashscope.aliyuncs.com/compatible-mode/v1'
                model = 'qwen-turbo'
            else:
                return '❌ API Key 未配置，请在系统设置中配置硅基流动或阿里云 API Key'
        else:
            base_url = 'https://api.siliconflow.cn/v1'
            model = 'Qwen/Qwen2.5-14B-Instruct'

        from openai import OpenAI
        client = OpenAI(api_key=api_key, base_url=base_url)
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {'role': 'system', 'content': _('You are a professional document writing assistant. Output only Markdown content, without any additional explanations.')},
                {'role': 'user', 'content': prompt_text}
            ],
            temperature=0.5,
            max_tokens=4096
        )
        content = resp.choices[0].message.content or ''
        if not content.strip():
            return '❌ AI 内容生成为空，请重试'

        # 保存为 .md 文件
        import uuid
        media_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'data', 'media', 'temp')
        os.makedirs(media_dir, exist_ok=True)
        fn = f"md_{uuid.uuid4().hex[:8]}.md"
        fp = os.path.join(media_dir, fn)
        with open(fp, 'w', encoding='utf-8') as f:
            f.write(content)

        download_url = f'/admin/agent-matrix/media/download/{fn}'
        preview_url = f'/admin/agent-matrix/md-preview/{fn}'
        preview = content[:200].strip().replace('\n', ' ')
        return f'✅ 文档已生成："{topic}"\n预览：{preview}...\n文件：{fn}\n预览链接：{preview_url}\n下载链接：{download_url}'

    except Exception as e:
        logger.warning(f"[tool:generate_markdown] 执行失败: {e}")
        return f'❌ Markdown Generation Error: {e}'


def _tool_generate_docx(args):
    """使用 AI 生成 Word 文档（.docx）"""
    try:
        topic = str(args.get('topic', '')).strip()
        if not topic:
            return '❌ 请提供文档主题'
        sections = max(3, min(int(args.get('sections', 5) or 5), 12))
        style = str(args.get('style', 'professional'))

        style_desc = {
            'professional': _('Professional business style, using formal titles and paragraph formats'),
            'formal': _('Formal official document style, clear chapters and sections, rigorous language'),
            'creative': _('Creative Design Style, Flexible Layouts and Modern Feel'),
            'simple': _('Concise style, key points highlighted for quick reading'),
            'report': _('Report Style: Includes summary, data analysis, conclusions, and recommendations'),
        }.get(style, _('Professional business style'))

        prompt_text = f'请撰写一篇关于"{topic}"的Word文档，包含{sections}个章节。\n风格要求：{style_desc}\n请输出Markdown格式，包含标题（##）、段落、列表。'

        # 读取 API Key（硅基流动优先）
        from models import get_db
        with get_db() as conn:
            row = conn.execute("SELECT value FROM system_config WHERE key='siliconflow_api_key'").fetchone()
        api_key = row['value'] if row else os.environ.get('SILICONFLOW_API_KEY', '')
        if not api_key:
            with get_db() as conn:
                row = conn.execute("SELECT value FROM system_config WHERE key='dashscope_text_key'").fetchone()
            if row:
                api_key = row['value']
                base_url, model = 'https://dashscope.aliyuncs.com/compatible-mode/v1', 'qwen-turbo'
            else:
                return _('❌ API Key Not Configured')
        else:
            base_url, model = 'https://api.siliconflow.cn/v1', 'Qwen/Qwen2.5-14B-Instruct'

        from openai import OpenAI
        client = OpenAI(api_key=api_key, base_url=base_url)
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {'role': 'system', 'content': _('You are a professional document writing assistant. Output structured Markdown content with multiple levels of headings and paragraphs, without any additional explanations.')},
                {'role': 'user', 'content': prompt_text}
            ],
            temperature=0.5, max_tokens=4096
        )
        md_content = resp.choices[0].message.content or ''
        if not md_content.strip():
            return _('❌ AI Content Generation is Empty')

        # 用 python-docx 渲染
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()

        # 设置默认字体
        style_n = doc.styles['Normal']
        style_n.font.name = 'Microsoft YaHei'
        style_n.font.size = Pt(11)
        style_n.paragraph_format.space_after = Pt(6)

        # 标题
        lines = md_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('## '):
                h = doc.add_heading(line[3:], level=2)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                    run.font.size = Pt(16)
            elif line.startswith('### '):
                h = doc.add_heading(line[4:], level=3)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x44)
                    run.font.size = Pt(13)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or line.startswith('2. ') or any(line.startswith(f'{i}. ') for i in range(1, 10)):
                p = doc.add_paragraph(line, style='List Number')
            else:
                p = doc.add_paragraph(line)

        import uuid
        media_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'data', 'media', 'temp')
        os.makedirs(media_dir, exist_ok=True)
        fn = f"doc_{uuid.uuid4().hex[:8]}.docx"
        fp = os.path.join(media_dir, fn)
        doc.save(fp)

        download_url = f'/admin/agent-matrix/media/download/{fn}'
        return f'✅ Word 文档已生成："{topic}"（{sections}章）\n文件：{fn}\n下载链接：{download_url}'

    except Exception as e:
        logger.warning(f"[tool:generate_docx] 执行失败: {e}")
        return f'❌ Word Document Generation Error: {e}'


TOOL_EXECUTORS = {
    "get_system_health": _tool_get_system_health,
    "query_stats": _tool_query_stats,
    "search_knowledge": _tool_search_knowledge,
    "ads_list": _tool_ads_list,
    "ads_create": _tool_ads_create,
    "ads_update": _tool_ads_update,
    "ads_delete": _tool_ads_delete,
    "ads_get_stats": _tool_ads_get_stats,
    "ads_analyze": _tool_ads_analyze,
    "ads_render_snippet": _tool_ads_render_snippet,
    "cms_list_channels": _tool_cms_list_channels,
    "cms_create_post": _tool_cms_create_post,
    "cms_publish_post": _tool_cms_publish_post,
    "generate_ppt": _tool_generate_ppt,
    "generate_image": _tool_generate_image,
    "generate_markdown": _tool_generate_markdown,
    "generate_docx": _tool_generate_docx,
}


# ============================================================
# 对外接口
# ============================================================

def get_tools_for_agent(allowed_tools):
    """按 Agent 的 allowed_tools 白名单返回可用工具 schema 列表。

    allowed_tools 可为 JSON 字符串或 list；为空/无效时返回空列表
    （即该 Agent 不启用工具，走原单轮逻辑）。
    """
    if not allowed_tools:
        return []
    if isinstance(allowed_tools, str):
        try:
            allowed_tools = json.loads(allowed_tools)
        except (json.JSONDecodeError, TypeError):
            return []
    if not isinstance(allowed_tools, list):
        return []
    return [TOOL_SCHEMAS[name] for name in allowed_tools if name in TOOL_SCHEMAS]


def execute_tool(name, args):
    """执行指定工具，返回字符串结果。未知工具或异常均返回错误字符串。"""
    executor = TOOL_EXECUTORS.get(name)
    if not executor:
        return f"Unknown tool: {name}"
    if not isinstance(args, dict):
        args = {}
    try:
        return executor(args)
    except Exception as e:
        logger.warning(f"[tool:{name}] 未捕获异常: {e}")
        return f"Tool {name} execution error: {e}"
